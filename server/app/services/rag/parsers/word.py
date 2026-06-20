"""Word 文档文本提取：.docx（python-docx）、.doc（doc2txt / antiword）。"""

from pathlib import Path

from app.services.rag.errors import IngestError


def _require_non_empty(text: str, *, empty_message: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise IngestError(empty_message)
    return cleaned


def parse_docx_file(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestError("DOCX 解析依赖未安装，请执行 pip install python-docx") from exc

    try:
        document = Document(str(file_path))
    except Exception as exc:
        raise IngestError(f"无法读取 DOCX: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return _require_non_empty("\n\n".join(parts), empty_message="DOCX 未提取到文本")


def parse_doc_file(file_path: Path) -> str:
    try:
        from doc2txt import extract_text
    except ImportError as exc:
        raise IngestError("DOC 解析依赖未安装，请执行 pip install doc2txt") from exc

    try:
        text = extract_text(str(file_path), optimize_format=True)
    except FileNotFoundError as exc:
        raise IngestError("DOC 文件不存在") from exc
    except ValueError as exc:
        raise IngestError(str(exc)) from exc
    except RuntimeError as exc:
        raise IngestError(f"无法读取 DOC: {exc}") from exc
    except Exception as exc:
        raise IngestError(f"无法读取 DOC: {exc}") from exc

    return _require_non_empty(text or "", empty_message="DOC 未提取到文本")
